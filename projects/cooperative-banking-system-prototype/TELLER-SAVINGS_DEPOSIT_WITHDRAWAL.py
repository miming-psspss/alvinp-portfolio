import sqlite3
import tkinter as tk
import sys
import os
from tkinter import ttk, messagebox
from datetime import datetime, timedelta
from decimal import Decimal, getcontext

# Set precision for decimal calculations
getcontext().prec = 10


class TellerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Teller Application")
        self.root.geometry("1200x800")
        
        # Database connections
        self.conn = sqlite3.connect("financial_records.db")
        self.savings_ledger = sqlite3.connect("savings_ledger.db")
        self.create_tables()
        
        # Transaction counter
        self.current_transaction_number = self.get_last_transaction_number() + 1
        
        # Cash on hand
        self.cash_on_hand = Decimal('0.0')
        self.load_cash_on_hand()
        
        # Create notebook (tabs)
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True)
        
        # Create tabs
        self.create_cash_management_tab()
        self.create_savings_transactions_tab()
        self.create_cheque_processing_tab()
        self.create_daily_report_tab()
        self.create_ledger_viewer_tab()
        self.create_interest_calculation_tab()
        
        # Load initial data
        self.update_cash_display()
    
    def get_last_transaction_number(self):
        """Get the last transaction number from the database"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT MAX(CAST(transaction_number AS INTEGER)) FROM teller_transactions")
            result = cursor.fetchone()
            return result[0] if result[0] is not None else 0
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error getting last transaction number: {e}")
            return 0
    
    def load_cash_on_hand(self):
        """Load the current cash on hand balance"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT balance FROM teller_cash ORDER BY date DESC LIMIT 1")
            result = cursor.fetchone()
            if result:
                self.cash_on_hand = Decimal(str(result[0]))
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error loading cash balance: {e}")
            self.cash_on_hand = Decimal('0.0')
    
    def create_tables(self):
        """Create database tables if they don't exist with updated schema for interest calculation"""
        try:
            cursor = self.conn.cursor()
            
            # Main accounts table with interest calculation fields
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS accounts (
                    account_number TEXT PRIMARY KEY,
                    last_name TEXT NOT NULL,
                    first_name TEXT NOT NULL,
                    total_savings REAL NOT NULL,
                    lq_total_savings REAL NOT NULL DEFAULT 0,
                    savings_quarterly_interest REAL DEFAULT 0
                )
            """)
            
            # Teller transactions table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS teller_transactions (
                    transaction_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    transaction_number TEXT NOT NULL,
                    transaction_date TEXT NOT NULL,
                    transaction_type TEXT NOT NULL,
                    account_number TEXT NOT NULL,
                    amount REAL NOT NULL,
                    cash_denominations TEXT,
                    cheque_details TEXT,
                    teller_id TEXT NOT NULL
                )
            """)
            
            # Teller cash balance table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS teller_cash (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    date TEXT NOT NULL,
                    balance REAL NOT NULL
                )
            """)
            
            # Ledger connection tables
            ledger_cursor = self.savings_ledger.cursor()
            ledger_cursor.execute("""
                CREATE TABLE IF NOT EXISTS savings_ledger (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    transaction_id TEXT NOT NULL,
                    account_number TEXT NOT NULL,
                    transaction_date TEXT NOT NULL,
                    transaction_type TEXT NOT NULL,
                    amount REAL NOT NULL,
                    balance REAL NOT NULL,
                    description TEXT,
                    teller_id TEXT NOT NULL
                )
            """)
            
            ledger_cursor.execute("""
                CREATE TABLE IF NOT EXISTS account_balances (
                    account_number TEXT PRIMARY KEY,
                    last_updated TEXT NOT NULL,
                    current_balance REAL NOT NULL
                )
            """)
            
            # Passbook table for interest transactions
            ledger_cursor.execute("""
                CREATE TABLE IF NOT EXISTS passbook (
                    entry_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    account_number TEXT NOT NULL,
                    date TEXT NOT NULL,
                    description TEXT NOT NULL,
                    credit REAL DEFAULT 0,
                    debit REAL DEFAULT 0,
                    balance REAL NOT NULL,
                    FOREIGN KEY(account_number) REFERENCES accounts(account_number)
                )
            """)
            
            self.conn.commit()
            self.savings_ledger.commit()
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error creating tables: {e}")
    
    def find_account(self):
        """Find and display account information for the savings account and show ledger"""
        account_num = self.account_number.get().strip().zfill(4)
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Input Error", "Please enter a valid 4-digit account number")
            return
        
        try:
            # Clear existing ledger data
            for item in self.savings_ledger_tree.get_children():
                self.savings_ledger_tree.delete(item)
                
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT account_number, last_name, first_name, total_savings 
                FROM accounts 
                WHERE account_number = ?
            """, (account_num,))
            
            account = cursor.fetchone()
            if not account:
                messagebox.showerror("Error", "Account not found")
                self.account_info.config(text="")
                return
            
            self.account_info.config(
                text=f"Account: {account[0]} - {account[1]}, {account[2]}\nCurrent Savings: PHP{Decimal(account[3]):,.2f}"
            )
            
            # Get ledger entries from savings_ledger.db
            ledger_cursor = self.savings_ledger.cursor()
            ledger_cursor.execute("""
                SELECT transaction_date, transaction_type, amount, balance, description
                FROM savings_ledger
                WHERE account_number = ?
                ORDER BY transaction_date DESC
            """, (account_num,))
            
            # Add ledger entries to treeview
            for entry in ledger_cursor.fetchall():
                amount = Decimal(str(entry[2]))
                formatted_amount = f"PHP{amount:,.2f}" if amount >= 0 else f"(PHP{abs(amount):,.2f})"
                formatted_balance = f"PHP{Decimal(str(entry[3])):,.2f}"
                
                self.savings_ledger_tree.insert("", "end", values=(
                    entry[0],  # date
                    entry[1],  # type
                    formatted_amount,
                    formatted_balance,
                    entry[4]   # description
                ))
                
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error finding account: {e}")
            self.account_info.config(text="")
    
    def create_cash_management_tab(self):
        """Create the cash management tab with denomination inputs"""
        self.cash_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.cash_tab, text="Cash Management")
        
        # Cash on hand frame
        cash_frame = ttk.LabelFrame(self.cash_tab, text="Cash on Hand", padding=10)
        cash_frame.pack(fill='x', padx=10, pady=5)
        
        self.cash_label = ttk.Label(cash_frame, text="Current Cash: PHP0.00", font=('Arial', 14))
        self.cash_label.pack()
        
        # Transaction frame
        trans_frame = ttk.LabelFrame(self.cash_tab, text="Cash Transaction", padding=10)
        trans_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Label(trans_frame, text="Transaction Type:").grid(row=0, column=0, sticky='e')
        self.cash_trans_type = ttk.Combobox(trans_frame, 
                                          values=["Credit (Deposit)", "Debit (Withdrawal)"], 
                                          state="readonly")
        self.cash_trans_type.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        
        # Add denomination frame with proper variable initialization
        self.cash_denom_frame = ttk.LabelFrame(trans_frame, text="Cash Denominations", padding=10)
        self.cash_denom_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky='ew')
        
        # Initialize cash denominations
        self.cash_denominations = {
            '1000': {'label': 'PHP1000:', 'var': tk.StringVar(value='0'), 'row': 0},
            '500': {'label': 'PHP500:', 'var': tk.StringVar(value='0'), 'row': 1},
            '200': {'label': 'PHP200:', 'var': tk.StringVar(value='0'), 'row': 2},
            '100': {'label': 'PHP100:', 'var': tk.StringVar(value='0'), 'row': 3},
            '50': {'label': 'PHP50:', 'var': tk.StringVar(value='0'), 'row': 4},
            '20': {'label': 'PHP20:', 'var': tk.StringVar(value='0'), 'row': 5},
            '10': {'label': 'PHP10:', 'var': tk.StringVar(value='0'), 'row': 6},
            '5': {'label': 'PHP5:', 'var': tk.StringVar(value='0'), 'row': 7},
            '1': {'label': 'PHP1:', 'var': tk.StringVar(value='0'), 'row': 8}
        }
        
        # Create denomination inputs with validation
        for denom, data in self.cash_denominations.items():
            ttk.Label(self.cash_denom_frame, text=data['label']).grid(row=data['row'], column=0, sticky='e')
            entry = ttk.Entry(self.cash_denom_frame, textvariable=data['var'], width=10,
                             validate="key",
                             validatecommand=(trans_frame.register(self.validate_denomination), '%P'))
            entry.grid(row=data['row'], column=1, padx=5, pady=2, sticky='w')
            data['var'].trace_add('write', lambda *args, d=denom: self.update_cash_total())
        
        # Total amount display
        ttk.Label(self.cash_denom_frame, text="Total:").grid(row=9, column=0, sticky='e')
        self.cash_total_var = tk.StringVar(value="PHP0.00")
        ttk.Label(self.cash_denom_frame, textvariable=self.cash_total_var,
                 font=('Arial', 10, 'bold')).grid(row=9, column=1, sticky='w')
        
        ttk.Button(trans_frame, text="Process Transaction", 
                  command=self.process_cash_transaction).grid(row=2, column=1, pady=10)

    def create_savings_transactions_tab(self):
        """Create the savings transactions tab with ledger display"""
        self.savings_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.savings_tab, text="Savings Transactions")
        
        # Main container frame
        main_frame = ttk.Frame(self.savings_tab)
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Top section - Account information and transaction input
        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill='x', pady=5)
        
        # Account selection frame
        account_frame = ttk.LabelFrame(top_frame, text="Account Information", padding=10)
        account_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(account_frame, text="Account Number:").grid(row=0, column=0, sticky='e')
        self.account_number = ttk.Entry(account_frame)
        self.account_number.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Button(account_frame, text="Find Account", command=self.find_account).grid(row=0, column=2, padx=5)
        
        self.account_info = ttk.Label(account_frame, text="")
        self.account_info.grid(row=1, column=0, columnspan=3, pady=5)
        
        # Transaction frame
        trans_frame = ttk.LabelFrame(top_frame, text="Savings Transaction", padding=10)
        trans_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(trans_frame, text="Transaction Type:").grid(row=0, column=0, sticky='e')
        self.savings_trans_type = ttk.Combobox(trans_frame, values=["Deposit", "Withdrawal"], state="readonly")
        self.savings_trans_type.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        
        # Add denomination frame for savings
        self.savings_denom_frame = ttk.LabelFrame(trans_frame, text="Cash Denominations", padding=10)
        self.savings_denom_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky='ew')
        
        # Initialize savings denominations
        self.savings_denominations = {
            '1000': {'label': 'PHP1000:', 'var': tk.StringVar(value='0'), 'row': 0},
            '500': {'label': 'PHP500:', 'var': tk.StringVar(value='0'), 'row': 1},
            '200': {'label': 'PHP200:', 'var': tk.StringVar(value='0'), 'row': 2},
            '100': {'label': 'PHP100:', 'var': tk.StringVar(value='0'), 'row': 3},
            '50': {'label': 'PHP50:', 'var': tk.StringVar(value='0'), 'row': 4},
            '20': {'label': 'PHP20:', 'var': tk.StringVar(value='0'), 'row': 5},
            '10': {'label': 'PHP10:', 'var': tk.StringVar(value='0'), 'row': 6},
            '5': {'label': 'PHP5:', 'var': tk.StringVar(value='0'), 'row': 7},
            '1': {'label': 'PHP1:', 'var': tk.StringVar(value='0'), 'row': 8}
        }
        
        # Create denomination inputs with validation
        for denom, data in self.savings_denominations.items():
            ttk.Label(self.savings_denom_frame, text=data['label']).grid(row=data['row'], column=0, sticky='e')
            entry = ttk.Entry(self.savings_denom_frame, textvariable=data['var'], width=10,
                            validate="key",
                            validatecommand=(trans_frame.register(self.validate_denomination), '%P'))
            entry.grid(row=data['row'], column=1, padx=5, pady=2, sticky='w')
            data['var'].trace_add('write', lambda *args, d=denom: self.update_savings_total())
        
        # Total amount display
        ttk.Label(self.savings_denom_frame, text="Total:").grid(row=9, column=0, sticky='e')
        self.savings_total_var = tk.StringVar(value="PHP0.00")
        ttk.Label(self.savings_denom_frame, textvariable=self.savings_total_var,
                font=('Arial', 10, 'bold')).grid(row=9, column=1, sticky='w')
        
        # Button frame for both Process Transaction and Print Passbook Entry
        button_frame = ttk.Frame(trans_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Process Transaction", 
                command=self.process_savings_transaction).pack(side='left', padx=5)
        
        ttk.Button(button_frame, text="Print Passbook Entry", 
                command=self.print_current_savings_transaction).pack(side='left', padx=5)
        
        # Bottom section - Ledger display
        bottom_frame = ttk.Frame(main_frame)
        bottom_frame.pack(fill='both', expand=True, pady=5)
        
        # Treeview for ledger entries
        columns = ("date", "type", "amount", "balance", "description")
        self.savings_ledger_tree = ttk.Treeview(bottom_frame, columns=columns, show='headings')
        
        # Define headings
        self.savings_ledger_tree.heading("date", text="Date")
        self.savings_ledger_tree.heading("type", text="Type")
        self.savings_ledger_tree.heading("amount", text="Amount")
        self.savings_ledger_tree.heading("balance", text="Balance")
        self.savings_ledger_tree.heading("description", text="Description")
        
        # Set column widths
        self.savings_ledger_tree.column("date", width=120)
        self.savings_ledger_tree.column("type", width=100)
        self.savings_ledger_tree.column("amount", width=100)
        self.savings_ledger_tree.column("balance", width=100)
        self.savings_ledger_tree.column("description", width=200)
        
        # Add scrollbars
        scroll_y = ttk.Scrollbar(bottom_frame, orient='vertical', command=self.savings_ledger_tree.yview)
        scroll_x = ttk.Scrollbar(bottom_frame, orient='horizontal', command=self.savings_ledger_tree.xview)
        self.savings_ledger_tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Layout
        self.savings_ledger_tree.grid(row=0, column=0, sticky='nsew')
        scroll_y.grid(row=0, column=1, sticky='ns')
        scroll_x.grid(row=1, column=0, sticky='ew')
        bottom_frame.grid_rowconfigure(0, weight=1)
        bottom_frame.grid_columnconfigure(0, weight=1)
    
    def print_current_savings_transaction(self):
        """Print the current savings transaction in passbook format with updated headers"""
        account_num = self.account_number.get().strip().zfill(4)
        trans_type = self.savings_trans_type.get()
        total_amount = Decimal(str(self.update_savings_total()))
        
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Input Error", "Please enter a valid account number")
            return
        
        if not trans_type:
            messagebox.showwarning("Input Error", "Please select transaction type")
            return
        
        if total_amount <= 0:
            messagebox.showwarning("Input Error", "Please enter a positive amount")
            return
        
        try:
            # Get account info
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT last_name, first_name, total_savings 
                FROM accounts 
                WHERE account_number = ?
            """, (account_num,))
            account_info = cursor.fetchone()
            
            if not account_info:
                messagebox.showerror("Error", "Account not found")
                return
            
            last_name, first_name, current_balance = account_info
            
            # Calculate new balance for display
            if trans_type == "Deposit":
                new_balance = Decimal(str(current_balance)) + total_amount
                transaction_display = ""  # Deposits show amount in deposit column
                deposit_display = f"PHP{total_amount:,.2f}"
            else:
                new_balance = Decimal(str(current_balance)) - total_amount
                transaction_display = f"PHP{total_amount:,.2f}"  # Withdrawals show amount here
                deposit_display = ""
            
            # Get denomination string
            denominations = self.get_denomination_string(is_cash_tab=False)
            
            # Transaction type to code mapping
            type_codes = {
                "Deposit": "DEP",
                "Withdrawal": "WDL",
                "Cheque Deposit": "CHQ",
                "Interest Credit": "INT"
            }
            code = type_codes.get(trans_type, "OTH")
            
            # Create passbook-style report with updated headers
            report_lines = []
            report_lines.append("=" * 70)
            report_lines.append(f"PASSBOOK ENTRY - {datetime.now().strftime('%d/%m/%Y %H:%M')}".center(70))
            report_lines.append(f"ACCOUNT: {account_num} - {last_name}, {first_name}".center(70))
            report_lines.append("=" * 70)
            
            # Updated headers with consistent spacing
            report_lines.append(
                "DATE".ljust(12) + 
                "WITHDRAWAL".ljust(15) + 
                "DEPOSIT".ljust(15) + 
                "CODE".ljust(8) + 
                "BALANCE".ljust(15) + 
                "REF NO.".ljust(10)
            )
            
            report_lines.append("-" * 70)
            
            # Add the transaction with proper columns
            report_lines.append(
                datetime.now().strftime("%d/%m/%y").ljust(12) +
                transaction_display.ljust(15) +
                deposit_display.ljust(15) +
                code.ljust(8) +
                f"PHP{new_balance:,.2f}".ljust(15) +
                "NEW".ljust(10)  # Using "NEW" as reference for current transaction
            )
            
            report_lines.append("-" * 70)
            report_lines.append(f"Transaction Details: {trans_type} - {denominations}".center(70))
            report_lines.append("=" * 70)
            
            full_report = "\n".join(report_lines)
            
            # Show print dialog (reusing the method from Account Ledger tab)
            self.show_passbook_print_dialog(full_report, account_num)
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error generating passbook entry: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {str(e)}")

    def validate_denomination(self, new_value):
        """Validate denomination input"""
        if new_value == "":
            return True
        return new_value.isdigit() and int(new_value) >= 0

    def update_cash_total(self):
        """Calculate and update the total for cash transactions"""
        total = Decimal('0.0')
        for denom, data in self.cash_denominations.items():
            try:
                value = data['var'].get()
                count = int(value) if value else 0
                if count < 0:
                    count = 0
                    data['var'].set('0')
                total += Decimal(denom) * Decimal(count)
            except ValueError:
                data['var'].set('0')
                total += Decimal('0.0')
        
        self.cash_total_var.set(f"PHP{total:,.2f}")
        return float(total)

    def update_savings_total(self):
        """Calculate and update the total for savings transactions"""
        total = Decimal('0.0')
        for denom, data in self.savings_denominations.items():
            try:
                value = data['var'].get()
                count = int(value) if value else 0
                if count < 0:
                    count = 0
                    data['var'].set('0')
                total += Decimal(denom) * Decimal(count)
            except ValueError:
                data['var'].set('0')
                total += Decimal('0.0')
        
        self.savings_total_var.set(f"PHP{total:,.2f}")
        return float(total)

    def get_denomination_string(self, is_cash_tab=True):
        """Generate a string representation of denominations"""
        denom_list = []
        denominations = self.cash_denominations if is_cash_tab else self.savings_denominations
        
        for denom, data in denominations.items():
            try:
                count = int(data['var'].get() or 0)
                if count > 0:
                    denom_list.append(f"{data['label']} x{count}")
            except ValueError:
                continue
        
        return ", ".join(denom_list) if denom_list else "No denominations specified"

    def generate_transaction_number(self):
        """Generate a transaction number with date prefix"""
        today = datetime.now().strftime("%Y%m%d")
        trans_num = f"{today}-{self.current_transaction_number:04d}"
        self.current_transaction_number += 1
        return trans_num

    def update_cash_display(self):
        """Update the cash on hand display"""
        self.cash_label.config(text=f"Current Cash: PHP{self.cash_on_hand:,.2f}")

    def process_cash_transaction(self):
        """Process a cash transaction with denomination tracking"""
        trans_type = self.cash_trans_type.get()
        total_amount = Decimal(str(self.update_cash_total()))
        
        if not trans_type:
            messagebox.showwarning("Input Error", "Please select transaction type")
            return
        
        if total_amount <= 0:
            messagebox.showwarning("Input Error", "Please enter positive denomination values")
            return
        
        # Generate denomination string
        denominations = self.get_denomination_string(is_cash_tab=True)
        
        # Generate transaction number
        trans_num = self.generate_transaction_number()
        today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Update cash on hand
        if "Credit" in trans_type:
            self.cash_on_hand += total_amount
        else:
            if total_amount > self.cash_on_hand:
                messagebox.showerror("Error", "Insufficient cash on hand")
                return
            self.cash_on_hand -= total_amount
        
        try:
            # Save transaction
            cursor = self.conn.cursor()
            cursor.execute("""
                INSERT INTO teller_transactions 
                (transaction_number, transaction_date, transaction_type, 
                 account_number, amount, cash_denominations, teller_id)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                trans_num, today, trans_type, 
                "CASH", float(total_amount), denominations, "TELLER1"
            ))
            
            # Update cash balance
            cursor.execute("INSERT INTO teller_cash (date, balance) VALUES (?, ?)", 
                          (today, float(self.cash_on_hand)))
            
            self.conn.commit()
            self.update_cash_display()
            messagebox.showinfo("Success", f"Transaction {trans_num} processed successfully")
            
            # Clear denomination inputs
            for data in self.cash_denominations.values():
                data['var'].set('0')
            
        except sqlite3.Error as e:
            self.conn.rollback()
            messagebox.showerror("Database Error", f"Error processing transaction: {e}")

    def process_savings_transaction(self):
        """Process savings transactions without double-counting"""
        account_num = self.account_number.get().strip().zfill(4)
        trans_type = self.savings_trans_type.get()
        total_amount = Decimal(str(self.update_savings_total()))
        
        if not account_num or len(account_num) != 4 or not trans_type:
            messagebox.showwarning("Input Error", "Please fill all required fields with valid account number")
            return
        
        if total_amount <= 0:
            messagebox.showwarning("Input Error", "Please enter positive denomination values")
            return
        
        trans_num = self.generate_transaction_number()
        today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        denominations = self.get_denomination_string(is_cash_tab=False)
        
        try:
            # Start transaction
            self.conn.execute("BEGIN TRANSACTION")
            self.savings_ledger.execute("BEGIN TRANSACTION")
            
            # Get current balance (from either database)
            cursor = self.conn.cursor()
            cursor.execute("SELECT total_savings FROM accounts WHERE account_number = ?", (account_num,))
            account = cursor.fetchone()
            
            if not account:
                raise ValueError("Account not found")
                
            current_balance = Decimal(str(account[0]))
            
            # Calculate new balance (ONLY ONCE)
            if trans_type == "Deposit":
                new_balance = current_balance + total_amount
                # Update cash on hand
                self.cash_on_hand += total_amount
                cursor.execute("INSERT INTO teller_cash (date, balance) VALUES (?, ?)", 
                            (today, float(self.cash_on_hand)))
            else:
                if total_amount > current_balance:
                    raise ValueError("Insufficient account balance")
                new_balance = current_balance - total_amount
                if total_amount > self.cash_on_hand:
                    raise ValueError("Insufficient cash on hand")
                self.cash_on_hand -= total_amount
                cursor.execute("INSERT INTO teller_cash (date, balance) VALUES (?, ?)", 
                            (today, float(self.cash_on_hand)))
            
            # Update main accounts table (ONLY ONCE)
            cursor.execute("""
                UPDATE accounts 
                SET total_savings = ? 
                WHERE account_number = ?
            """, (float(new_balance), account_num))
            
            # Record transaction
            cursor.execute("""
                INSERT INTO teller_transactions 
                (transaction_number, transaction_date, transaction_type, 
                account_number, amount, cash_denominations, teller_id)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (trans_num, today, trans_type, account_num, 
                float(total_amount), denominations, "TELLER1"))
            
            # Update ledger (without changing the balance again)
            self.update_savings_ledger(
                transaction_id=trans_num,
                account_number=account_num,
                trans_type=trans_type,
                amount=float(total_amount),
                new_balance=float(new_balance),  # Pass the already-calculated balance
                description=f"{trans_type} - {denominations}"
            )
            
            # Commit transactions
            self.conn.commit()
            self.savings_ledger.commit()
            
            self.update_cash_display()
            messagebox.showinfo("Success", f"Transaction {trans_num} processed successfully")
            
            # Clear form
            for data in self.savings_denominations.values():
                data['var'].set('0')
            self.find_account()
            
        except Exception as e:
            self.conn.rollback()
            self.savings_ledger.rollback()
            messagebox.showerror("Error", str(e))

    def update_savings_ledger(self, transaction_id, account_number, trans_type, amount, new_balance, description):
        """Update ledger database with proper account initialization"""
        try:
            today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor = self.savings_ledger.cursor()
            
            # Check if account exists in ledger
            cursor.execute("SELECT 1 FROM account_balances WHERE account_number = ?", (account_number,))
            if not cursor.fetchone():
                # Initialize new account in ledger with current balance
                cursor.execute("""
                    INSERT INTO account_balances (account_number, last_updated, current_balance)
                    VALUES (?, ?, ?)
                """, (account_number, today, new_balance))
            
            # Insert ledger record
            cursor.execute("""
                INSERT INTO savings_ledger (
                    transaction_id, account_number, transaction_date, 
                    transaction_type, amount, balance, description, teller_id
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                transaction_id, account_number, today,
                trans_type, amount, new_balance, description, "TELLER1"
            ))
            
            # Update account balance in ledger
            cursor.execute("""
                UPDATE account_balances 
                SET current_balance = ?, last_updated = ?
                WHERE account_number = ?
            """, (new_balance, today, account_number))
            
            self.savings_ledger.commit()
            
        except Exception as e:
            self.savings_ledger.rollback()
            raise ValueError(f"Failed to update ledger: {str(e)}")
    
    def create_cheque_processing_tab(self):
        """Create the cheque processing tab with ledger display"""
        self.cheque_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.cheque_tab, text="Cheque Processing")
        
        # Main container frame
        main_frame = ttk.Frame(self.cheque_tab)
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Top section - Account information and cheque details
        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill='x', pady=5)
        
        # Account selection frame
        account_frame = ttk.LabelFrame(top_frame, text="Account Information", padding=10)
        account_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(account_frame, text="Account Number:").grid(row=0, column=0, sticky='e')
        self.cheque_account_number = ttk.Entry(account_frame)
        self.cheque_account_number.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Button(account_frame, text="Find Account", command=self.find_cheque_account).grid(row=0, column=2, padx=5)
        
        self.cheque_account_info = ttk.Label(account_frame, text="")
        self.cheque_account_info.grid(row=1, column=0, columnspan=3, pady=5)
        
        # Cheque details frame
        cheque_frame = ttk.LabelFrame(top_frame, text="Cheque Details", padding=10)
        cheque_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(cheque_frame, text="Issuing Bank:").grid(row=0, column=0, sticky='e')
        self.issuing_bank = ttk.Entry(cheque_frame)
        self.issuing_bank.grid(row=0, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Label(cheque_frame, text="Date Issued (YYYY-MM-DD):").grid(row=1, column=0, sticky='e')
        self.cheque_date = ttk.Entry(cheque_frame)
        self.cheque_date.grid(row=1, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Label(cheque_frame, text="Cheque Number:").grid(row=2, column=0, sticky='e')
        self.cheque_number = ttk.Entry(cheque_frame)
        self.cheque_number.grid(row=2, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Label(cheque_frame, text="Amount:").grid(row=3, column=0, sticky='e')
        self.cheque_amount = ttk.Entry(cheque_frame)
        self.cheque_amount.grid(row=3, column=1, padx=5, pady=5, sticky='w')
        
        ttk.Button(cheque_frame, text="Process Cheque", command=self.process_cheque).grid(row=4, column=1, pady=10)
        
        # Bottom section - Ledger display
        bottom_frame = ttk.Frame(main_frame)
        bottom_frame.pack(fill='both', expand=True, pady=5)
        
        # Treeview for ledger entries
        columns = ("date", "type", "amount", "balance", "description")
        self.cheque_ledger_tree = ttk.Treeview(bottom_frame, columns=columns, show='headings')
        
        # Define headings
        self.cheque_ledger_tree.heading("date", text="Date")
        self.cheque_ledger_tree.heading("type", text="Type")
        self.cheque_ledger_tree.heading("amount", text="Amount")
        self.cheque_ledger_tree.heading("balance", text="Balance")
        self.cheque_ledger_tree.heading("description", text="Description")
        
        # Set column widths
        self.cheque_ledger_tree.column("date", width=120)
        self.cheque_ledger_tree.column("type", width=100)
        self.cheque_ledger_tree.column("amount", width=100)
        self.cheque_ledger_tree.column("balance", width=100)
        self.cheque_ledger_tree.column("description", width=200)
        
        # Add scrollbars
        scroll_y = ttk.Scrollbar(bottom_frame, orient='vertical', command=self.cheque_ledger_tree.yview)
        scroll_x = ttk.Scrollbar(bottom_frame, orient='horizontal', command=self.cheque_ledger_tree.xview)
        self.cheque_ledger_tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Layout
        self.cheque_ledger_tree.grid(row=0, column=0, sticky='nsew')
        scroll_y.grid(row=0, column=1, sticky='ns')
        scroll_x.grid(row=1, column=0, sticky='ew')
        bottom_frame.grid_rowconfigure(0, weight=1)
        bottom_frame.grid_columnconfigure(0, weight=1)
    
    def find_cheque_account(self):
        """Find and display account information for cheque processing and show ledger"""
        account_num = self.cheque_account_number.get().strip().zfill(4)
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Input Error", "Please enter a valid 4-digit account number")
            return
        
        try:
            # Clear existing ledger data
            for item in self.cheque_ledger_tree.get_children():
                self.cheque_ledger_tree.delete(item)
                
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT account_number, last_name, first_name 
                FROM accounts 
                WHERE account_number = ?
            """, (account_num,))
            
            account = cursor.fetchone()
            if not account:
                messagebox.showerror("Error", "Account not found")
                self.cheque_account_info.config(text="")
                return
            
            self.cheque_account_info.config(
                text=f"Account: {account[0]} - {account[1]}, {account[2]}"
            )
            
            # Get ledger entries from savings_ledger.db
            ledger_cursor = self.savings_ledger.cursor()
            ledger_cursor.execute("""
                SELECT transaction_date, transaction_type, amount, balance, description
                FROM savings_ledger
                WHERE account_number = ?
                ORDER BY transaction_date DESC
            """, (account_num,))
            
            # Add ledger entries to treeview
            for entry in ledger_cursor.fetchall():
                amount = Decimal(str(entry[2]))
                formatted_amount = f"PHP{amount:,.2f}" if amount >= 0 else f"(PHP{abs(amount):,.2f})"
                formatted_balance = f"PHP{Decimal(str(entry[3])):,.2f}"
                
                self.cheque_ledger_tree.insert("", "end", values=(
                    entry[0],  # date
                    entry[1],  # type
                    formatted_amount,
                    formatted_balance,
                    entry[4]   # description
                ))
                
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error finding account: {e}")
            self.cheque_account_info.config(text="")
    
    def validate_date(self, date_text):
        """Validate date format (YYYY-MM-DD)"""
        try:
            datetime.strptime(date_text, "%Y-%m-%d")
            return True
        except ValueError:
            return False
    
    def process_cheque(self):
        """Process a cheque deposit with proper balance updating in both databases"""
        account_num = self.cheque_account_number.get().strip().zfill(4)
        issuing_bank = self.issuing_bank.get().strip()
        cheque_date = self.cheque_date.get().strip()
        cheque_number = self.cheque_number.get().strip()
        amount_str = self.cheque_amount.get().strip()

        # Validate inputs
        if not all([account_num, issuing_bank, cheque_date, cheque_number, amount_str]):
            messagebox.showwarning("Input Error", "Please fill all required fields")
            return

        if len(account_num) != 4:
            messagebox.showwarning("Input Error", "Account number must be 4 digits")
            return

        if not self.validate_date(cheque_date):
            messagebox.showwarning("Input Error", "Invalid date format. Please use YYYY-MM-DD")
            return

        try:
            amount = Decimal(amount_str)
            if amount <= 0:
                raise ValueError("Amount must be positive")
        except ValueError:
            messagebox.showerror("Input Error", "Please enter a valid positive amount")
            return

        try:
            # Start transactions
            self.conn.execute("BEGIN TRANSACTION")
            self.savings_ledger.execute("BEGIN TRANSACTION")

            # Get current balance from main database
            cursor = self.conn.cursor()
            cursor.execute("SELECT total_savings FROM accounts WHERE account_number = ?", (account_num,))
            result = cursor.fetchone()
            
            if not result:
                raise ValueError("Account not found")
            
            current_balance = Decimal(str(result[0]))
            new_balance = current_balance + amount

            # Create cheque details string
            cheque_details = (
                f"Issuing Bank: {issuing_bank}, "
                f"Cheque #: {cheque_number}, "
                f"Date: {cheque_date}, "
                f"Amount: PHP{amount:,.2f}"
            )

            # Generate transaction number
            trans_num = self.generate_transaction_number()
            today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # Update main account balance
            cursor.execute("""
                UPDATE accounts 
                SET total_savings = ? 
                WHERE account_number = ?
            """, (float(new_balance), account_num))

            # Record transaction
            cursor.execute("""
                INSERT INTO teller_transactions 
                (transaction_number, transaction_date, transaction_type, 
                account_number, amount, cheque_details, teller_id)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                trans_num, today, "Cheque Deposit", 
                account_num, float(amount), cheque_details, "TELLER1"
            ))

            # Update ledger with the new balance
            self.update_savings_ledger(
                transaction_id=trans_num,
                account_number=account_num,
                trans_type="Cheque Deposit",
                amount=float(amount),
                new_balance=float(new_balance),
                description=cheque_details
            )

            # Commit transactions
            self.conn.commit()
            self.savings_ledger.commit()

            messagebox.showinfo("Success", f"Cheque {cheque_number} processed successfully")

            # Clear form
            self.issuing_bank.delete(0, tk.END)
            self.cheque_date.delete(0, tk.END)
            self.cheque_number.delete(0, tk.END)
            self.cheque_amount.delete(0, tk.END)
            self.find_cheque_account()

        except Exception as e:
            self.conn.rollback()
            self.savings_ledger.rollback()
            messagebox.showerror("Error", f"Failed to process cheque: {str(e)}")
    
    def create_daily_report_tab(self):
        """Create the daily report tab with date selection and printing options"""
        self.report_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.report_tab, text="Daily Report")
        
        # Report controls frame
        controls_frame = ttk.LabelFrame(self.report_tab, text="Report Controls", padding=10)
        controls_frame.pack(fill='x', padx=10, pady=5)
        
        # Date selection
        ttk.Label(controls_frame, text="Report Date:").grid(row=0, column=0, padx=5, sticky='e')
        self.report_date = ttk.Entry(controls_frame)
        self.report_date.grid(row=0, column=1, padx=5, sticky='w')
        self.report_date.insert(0, datetime.now().strftime("%Y-%m-%d"))
        
        # Calendar button for date selection
        ttk.Button(controls_frame, text="📅", width=3, 
                command=self.show_date_picker).grid(row=0, column=2, padx=5)
        
        # Action buttons
        ttk.Button(controls_frame, text="Generate Report", 
                command=self.generate_daily_report).grid(row=0, column=3, padx=5)
        ttk.Button(controls_frame, text="Print Report", 
                command=self.print_report).grid(row=0, column=4, padx=5)
        ttk.Button(controls_frame, text="Export to CSV", 
                command=self.save_report_csv).grid(row=0, column=5, padx=5)
        
        # Report display frame
        report_frame = ttk.Frame(self.report_tab)
        report_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Treeview for displaying transactions
        columns = ("trans_num", "date", "type", "account", "amount", "details")
        self.report_tree = ttk.Treeview(report_frame, columns=columns, show='headings')
        
        # Define headings
        self.report_tree.heading("trans_num", text="Trans #")
        self.report_tree.heading("date", text="Date")
        self.report_tree.heading("type", text="Type")
        self.report_tree.heading("account", text="Account")
        self.report_tree.heading("amount", text="Amount")
        self.report_tree.heading("details", text="Details")
        
        # Set column widths
        self.report_tree.column("trans_num", width=80)
        self.report_tree.column("date", width=120)
        self.report_tree.column("type", width=100)
        self.report_tree.column("account", width=80)
        self.report_tree.column("amount", width=100)
        self.report_tree.column("details", width=200)
        
        # Add scrollbars
        scroll_y = ttk.Scrollbar(report_frame, orient='vertical', command=self.report_tree.yview)
        scroll_x = ttk.Scrollbar(report_frame, orient='horizontal', command=self.report_tree.xview)
        self.report_tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Layout
        self.report_tree.grid(row=0, column=0, sticky='nsew')
        scroll_y.grid(row=0, column=1, sticky='ns')
        scroll_x.grid(row=1, column=0, sticky='ew')
        report_frame.grid_rowconfigure(0, weight=1)
        report_frame.grid_columnconfigure(0, weight=1)

    def show_date_picker(self):
        """Show a calendar dialog for date selection"""
        from tkcalendar import Calendar
        
        def set_date():
            self.report_date.delete(0, tk.END)
            self.report_date.insert(0, cal.selection_get().strftime("%Y-%m-%d"))
            top.destroy()
        
        top = tk.Toplevel(self.root)
        cal = Calendar(top, selectmode='day', date_pattern='y-mm-dd')
        cal.pack(padx=10, pady=10)
        ttk.Button(top, text="OK", command=set_date).pack(pady=5)

    def generate_daily_report(self):
        """Generate a report for the selected date"""
        report_date = self.report_date.get().strip()
        
        if not report_date:
            messagebox.showwarning("Input Error", "Please select a report date")
            return
        
        try:
            # Validate date format
            datetime.strptime(report_date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("Input Error", "Invalid date format. Please use YYYY-MM-DD")
            return
        
        # Clear existing data
        for item in self.report_tree.get_children():
            self.report_tree.delete(item)
        
        try:
            cursor = self.conn.cursor()
            
            # Get cash on hand as of report date
            cursor.execute("""
                SELECT balance 
                FROM teller_cash 
                WHERE date(date) <= date(?)
                ORDER BY date DESC
                LIMIT 1
            """, (report_date,))
            cash_result = cursor.fetchone()
            cash_on_hand = Decimal(str(cash_result[0])) if cash_result else Decimal('0.0')
            
            # Get transactions for the day
            cursor.execute("""
                SELECT transaction_number, transaction_date, transaction_type, 
                    account_number, amount, 
                    COALESCE(cash_denominations, cheque_details, '') as details
                FROM teller_transactions
                WHERE date(transaction_date) = date(?)
                ORDER BY transaction_date
            """, (report_date,))
            
            transactions = cursor.fetchall()
            
            if not transactions:
                messagebox.showinfo("Report", f"No transactions found for {report_date}")
                return
            
            # Add transactions to treeview
            for trans in transactions:
                self.report_tree.insert("", "end", values=(
                    trans[0],
                    trans[1],
                    trans[2],
                    trans[3],
                    f"PHP{float(trans[4]):,.2f}",
                    trans[5]
                ))
            
            # Calculate totals
            cursor.execute("""
                SELECT 
                    SUM(CASE WHEN transaction_type IN ('Credit (Deposit)', 'Deposit') THEN amount ELSE 0 END) as total_credits,
                    SUM(CASE WHEN transaction_type IN ('Debit (Withdrawal)', 'Withdrawal') THEN amount ELSE 0 END) as total_debits,
                    SUM(CASE WHEN transaction_type = 'Cheque Deposit' THEN amount ELSE 0 END) as total_cheques
                FROM teller_transactions
                WHERE date(transaction_date) = date(?)
            """, (report_date,))
            
            totals = cursor.fetchone()
            credits = totals[0] or 0
            debits = totals[1] or 0
            cheques = totals[2] or 0
            net_change = credits - debits
            
            # Add totals row with cash on hand
            self.report_tree.insert("", "end", values=(
                "",
                "TOTALS:",
                "",
                "",
                f"Credits: PHP{credits:,.2f}\nDebits: PHP{debits:,.2f}\nCheques: PHP{cheques:,.2f}",
                f"Net Change: PHP{net_change:,.2f}\nCash on Hand: PHP{cash_on_hand:,.2f}"
            ))
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error generating report: {e}")

    def print_report(self):
        """Generate TXT report with aligned columns and save to custom directory"""
        from datetime import datetime
        import os
        import platform

        report_date = self.report_date.get().strip()
        if not report_date:
            messagebox.showwarning("Input Error", "Please select a report date")
            return

        try:
            cursor = self.conn.cursor()
            
            # Get cash on hand as of report date
            cursor.execute("""
                SELECT balance 
                FROM teller_cash 
                WHERE date(date) <= date(?)
                ORDER BY date DESC
                LIMIT 1
            """, (report_date,))
            cash_result = cursor.fetchone()
            cash_on_hand = Decimal(str(cash_result[0])) if cash_result else Decimal('0.0')

            cursor.execute("""
                SELECT transaction_number, transaction_date, transaction_type, 
                    account_number, amount, 
                    COALESCE(cash_denominations, cheque_details, '') as details
                FROM teller_transactions
                WHERE date(transaction_date) = date(?)
                ORDER BY transaction_date
            """, (report_date,))
            transactions = cursor.fetchall()

            if not transactions:
                messagebox.showinfo("No Data", f"No transactions found for {report_date}")
                return

            cursor.execute("""
                SELECT 
                    SUM(CASE WHEN transaction_type IN ('Credit (Deposit)', 'Deposit') THEN amount ELSE 0 END),
                    SUM(CASE WHEN transaction_type IN ('Debit (Withdrawal)', 'Withdrawal') THEN amount ELSE 0 END),
                    SUM(CASE WHEN transaction_type = 'Cheque Deposit' THEN amount ELSE 0 END)
                FROM teller_transactions
                WHERE date(transaction_date) = date(?)
            """, (report_date,))
            totals = cursor.fetchone()
            credits = totals[0] or 0
            debits = totals[1] or 0
            cheques = totals[2] or 0
            net_change = credits - debits

            # Column setup
            headers = ["Trans#", "Date&Time", "Type", "Account", "Amount (PHP)", "Details"]
            col_widths = [20, 22, 20, 14, 16, 40]

            divider = "-" * sum(col_widths)
            header_row = "".join(h.ljust(col_widths[i]) for i, h in enumerate(headers))

            # Build content
            lines = []
            lines.append(f"Daily Teller Report - {report_date}")
            lines.append(divider)
            lines.append(header_row)
            lines.append(divider)

            for trans in transactions:
                amount = float(trans[4])
                row_cells = [
                    str(trans[0]),
                    str(trans[1]),
                    str(trans[2]),
                    str(trans[3]),
                    f"{amount:,.2f}",
                    str(trans[5])
                ]
                line = "".join(val.ljust(col_widths[i]) for i, val in enumerate(row_cells))
                lines.append(line)

            lines.append(divider)
            lines.append("")
            lines.append(f"{'Total Credits:':<25} {credits:>15,.2f} PHP")
            lines.append(f"{'Total Debits:':<25}  {debits:>15,.2f} PHP")
            lines.append(f"{'Total Cheques:':<25} {cheques:>15,.2f} PHP")
            lines.append(f"{'Net Change:':<25}   {net_change:>15,.2f} PHP")
            lines.append(f"{'Cash on Hand:':<25} {cash_on_hand:>15,.2f} PHP")
            lines.append(f"\nPrinted on {datetime.now().strftime('%Y-%m-%d %H:%M')}")

            # Custom save directory
            save_dir = r"C:\Users\PC\Documents\Python\FINANCIAL BANKING SYSTEM\DATABASE"
            os.makedirs(save_dir, exist_ok=True)  # Ensure the folder exists
            txt_path = os.path.join(save_dir, f"teller_report_{report_date}.txt")

            with open(txt_path, 'w', encoding='utf-8') as file:
                file.write('\n'.join(lines))

            # Open the file
            if platform.system() == 'Darwin':
                os.system(f'open "{txt_path}"')
            elif platform.system() == 'Windows':
                os.startfile(txt_path)
            else:
                os.system(f'xdg-open "{txt_path}"')

            messagebox.showinfo("Success", f"Report saved to:\n{txt_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{str(e)}")

    def save_report_csv(self):
        """Save the daily report to CSV, Excel, and Word files"""
        import csv, os, sys
        import xlsxwriter
        from docx import Document
        from tkinter import messagebox

        report_date = self.report_date.get().strip()
        if not report_date:
            messagebox.showwarning("Input Error", "Please select a report date")
            return

        filename_csv = f"daily_report_{report_date}.csv"
        filename_xlsx = f"daily_report_{report_date}.xlsx"
        filename_docx = f"daily_report_{report_date}.docx"

        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT transaction_number, transaction_date, transaction_type, 
                    account_number, amount, 
                    COALESCE(cash_denominations, cheque_details, '') as details
                FROM teller_transactions
                WHERE date(transaction_date) = date(?)
                ORDER BY transaction_date
            """, (report_date,))
            
            transactions = cursor.fetchall()
            if not transactions:
                messagebox.showwarning("Error", "No transactions to save")
                return

            headers = ["Transaction #", "Date", "Type", "Account", "Amount (PHP)", "Details"]

            # 💾 Write CSV
            with open(filename_csv, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                csvfile.write('\ufeff')  # Excel BOM
                writer.writerow(headers)
                for trans in transactions:
                    formatted_amount = f"PHP{float(trans[4]):,.2f}"
                    writer.writerow([trans[0], trans[1], trans[2], trans[3], formatted_amount, trans[5]])

            # 📊 Write Excel
            workbook = xlsxwriter.Workbook(filename_xlsx)
            sheet = workbook.add_worksheet("Daily Report")

            currency_fmt = workbook.add_format({'num_format': 'PHP#,##0.00', 'align': 'right'})
            bold_fmt = workbook.add_format({'bold': True})
            sheet.write_row(0, 0, headers, bold_fmt)

            for row_idx, trans in enumerate(transactions, start=1):
                sheet.write(row_idx, 0, trans[0])
                sheet.write(row_idx, 1, trans[1])
                sheet.write(row_idx, 2, trans[2])
                sheet.write(row_idx, 3, trans[3])
                sheet.write(row_idx, 4, float(trans[4]), currency_fmt)
                sheet.write(row_idx, 5, trans[5])

            workbook.close()

            # 📄 Write Word
            doc = Document()
            doc.add_heading(f"Daily Report – {report_date}", level=1)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            for idx, header in enumerate(headers):
                table.cell(0, idx).text = header

            for trans in transactions:
                row = table.add_row().cells
                row[0].text = str(trans[0])
                row[1].text = trans[1]
                row[2].text = trans[2]
                row[3].text = trans[3]
                row[4].text = f"PHP{float(trans[4]):,.2f}"
                row[5].text = trans[5]

            doc.save(filename_docx)

            messagebox.showinfo("Success", f"Reports saved:\n{filename_csv}\n{filename_xlsx}\n{filename_docx}")

            # 🔓 Open files in default apps (optional)
            if sys.platform == "win32":
                os.startfile(filename_xlsx)
                os.startfile(filename_docx)

        except PermissionError:
            messagebox.showerror("Error", "Please close the files before saving.")
        except Exception as e:
            messagebox.showerror("Error", f"Error saving report: {str(e)}")
    
    def create_ledger_viewer_tab(self):
        """Add a tab to view individual account ledgers with enhanced printing options"""
        self.ledger_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.ledger_tab, text="Account Ledger")
        
        # Account selection frame
        account_frame = ttk.Frame(self.ledger_tab, padding=10)
        account_frame.pack(fill='x')
        
        ttk.Label(account_frame, text="Account Number:").grid(row=0, column=0, padx=5, pady=5)
        self.ledger_account_entry = ttk.Entry(account_frame)
        self.ledger_account_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Button(account_frame, text="View Ledger", 
                command=self.view_account_ledger).grid(row=0, column=2, padx=5)
        
        # Add both print buttons
        ttk.Button(account_frame, text="Print Passbook", 
                command=self.print_account_passbook).grid(row=0, column=3, padx=5)
        ttk.Button(account_frame, text="Print Ledger", 
                command=self.print_account_ledger).grid(row=0, column=4, padx=5)
        
        # Treeview for ledger entries with additional columns
        columns = ("date", "type", "code", "ref_no", "amount", "balance", "description")
        self.ledger_tree = ttk.Treeview(self.ledger_tab, columns=columns, show='headings')
        
        # Define headings
        self.ledger_tree.heading("date", text="Date")
        self.ledger_tree.heading("type", text="Type")
        self.ledger_tree.heading("code", text="Code")
        self.ledger_tree.heading("ref_no", text="Ref No.")
        self.ledger_tree.heading("amount", text="Amount")
        self.ledger_tree.heading("balance", text="Balance")
        self.ledger_tree.heading("description", text="Description")
        
        # Set column widths
        self.ledger_tree.column("date", width=120)
        self.ledger_tree.column("type", width=100)
        self.ledger_tree.column("code", width=60)
        self.ledger_tree.column("ref_no", width=80)
        self.ledger_tree.column("amount", width=100)
        self.ledger_tree.column("balance", width=100)
        self.ledger_tree.column("description", width=200)
        
        # Add scrollbars
        scroll_y = ttk.Scrollbar(self.ledger_tab, orient='vertical', command=self.ledger_tree.yview)
        scroll_x = ttk.Scrollbar(self.ledger_tab, orient='horizontal', command=self.ledger_tree.xview)
        self.ledger_tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Layout
        self.ledger_tree.pack(side='left', fill='both', expand=True)
        scroll_y.pack(side='right', fill='y')
        scroll_x.pack(side='bottom', fill='x')

    def view_account_ledger(self):
        """Display the transaction ledger for an account with codes and reference numbers"""
        account_num = self.ledger_account_entry.get().strip().zfill(4)
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Error", "Please enter a valid 4-digit account number")
            return
        
        # Clear existing data
        for item in self.ledger_tree.get_children():
            self.ledger_tree.delete(item)
        
        try:
            cursor = self.savings_ledger.cursor()
            
            # Verify account exists
            cursor.execute("SELECT 1 FROM account_balances WHERE account_number = ?", (account_num,))
            if not cursor.fetchone():
                messagebox.showerror("Error", "Account not found in ledger")
                return
            
            # Get ledger entries with transaction_id as reference number
            cursor.execute("""
                SELECT transaction_date, transaction_type, transaction_id, 
                    amount, balance, description
                FROM savings_ledger
                WHERE account_number = ?
                ORDER BY transaction_date DESC
            """, (account_num,))
            
            # Transaction type to code mapping
            type_codes = {
                "Deposit": "DEP",
                "Withdrawal": "WDL",
                "Cheque Deposit": "CHQ",
                "Interest Credit": "INT",
                "Debit (Withdrawal)": "WDL",
                "Credit (Deposit)": "DEP"
            }
            
            for row in cursor.fetchall():
                # Determine transaction code
                trans_type = row[1]
                code = type_codes.get(trans_type, "OTH")
                
                # Format reference number - handle both string and numeric IDs
                try:
                    # First try to convert to integer if it's numeric
                    ref_no = f"{int(row[2]):05d}"  # 5-digit reference number with leading zeros
                except (ValueError, TypeError):
                    # If conversion fails, use the string as-is but pad with zeros
                    ref_no = row[2].zfill(5)
                    
                # Show withdrawals as negative amounts
                amount = Decimal(str(row[3]))
                if trans_type in ("Withdrawal", "Debit (Withdrawal)"):
                    amount = -amount
                    
                self.ledger_tree.insert("", "end", values=(
                    row[0],  # date
                    trans_type,  # type
                    code,  # code
                    ref_no,  # reference number
                    f"PHP{amount:,.2f}",  # amount (negative for withdrawals)
                    f"PHP{Decimal(str(row[4])):,.2f}",  # balance
                    row[5]  # description
                ))
            
            # Show current balance
            cursor.execute("SELECT current_balance FROM account_balances WHERE account_number = ?", (account_num,))
            balance = Decimal(str(cursor.fetchone()[0]))
            messagebox.showinfo("Account Status", f"Current Balance: PHP{balance:,.2f}")
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error viewing ledger: {e}")

    def print_account_passbook(self):
        account_num = self.ledger_account_entry.get().strip().zfill(4)
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Error", "Please enter a valid 4-digit account number")
            return
        
        try:
            # Use separate connections for each database
            main_cursor = self.conn.cursor()
            ledger_cursor = self.savings_ledger.cursor()
            
            # Get account info from main DB
            main_cursor.execute("""
                SELECT last_name, first_name 
                FROM accounts 
                WHERE account_number = ?
            """, (account_num,))
            account_info = main_cursor.fetchone()
            
            if not account_info:
                messagebox.showerror("Error", "Account not found")
                return
            
            last_name, first_name = account_info
            
            # Get balance from ledger DB
            ledger_cursor.execute("""
                SELECT current_balance 
                FROM account_balances 
                WHERE account_number = ?
            """, (account_num,))
            balance_result = ledger_cursor.fetchone()
            current_balance = Decimal(str(balance_result[0])) if balance_result else Decimal('0.0')
            
            # Get transactions from ledger DB
            ledger_cursor.execute("""
                SELECT transaction_date, transaction_type, transaction_id, 
                    amount, balance, description
                FROM savings_ledger
                WHERE account_number = ?
                ORDER BY transaction_date
            """, (account_num,))
            
            transactions = ledger_cursor.fetchall()
            
            if not transactions:
                messagebox.showwarning("Error", "No transactions found for this account")
                return
            
            # Create passbook report
            report_lines = []
            
            # Passbook header
            report_lines.append("=" * 60)
            report_lines.append(f"PASSBOOK FOR ACCOUNT: {account_num}".center(60))
            report_lines.append(f"NAME: {last_name}, {first_name}".center(60))
            report_lines.append("=" * 60)
            report_lines.append("DATE".ljust(15) + "W/D".ljust(10) + "CODE".ljust(10) + 
                            "BALANCE".ljust(15) + "REF NO.".ljust(10))
            report_lines.append("-" * 60)
            
            # Add transactions
            for trans in transactions:
                date = datetime.strptime(trans[0], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%y")
                trans_type = trans[1]
                
                # Handle reference number - ensure it's a string and pad with zeros
                ref_no = str(trans[2]).zfill(5)  # Convert to string and zero-pad
                
                amount = Decimal(str(trans[3]))
                balance = Decimal(str(trans[4]))
                
                # Transaction type to code mapping
                type_codes = {
                    "Deposit": "DEP",
                    "Withdrawal": "WDL",
                    "Cheque Deposit": "CHQ",
                    "Interest Credit": "INT",
                    "Debit (Withdrawal)": "WDL",
                    "Credit (Deposit)": "DEP"
                }
                code = type_codes.get(trans_type, "OTH")
                
                # Determine if withdrawal or deposit
                if trans_type in ("Withdrawal", "Debit (Withdrawal)"):
                    wd = f"PHP{amount:,.2f}"
                else:
                    wd = ""
                
                report_lines.append(
                    date.ljust(15) + 
                    wd.ljust(10) + 
                    code.ljust(10) + 
                    f"PHP{balance:,.2f}".ljust(15) + 
                    ref_no.ljust(10)
                )
            
            # Add footer with current balance
            report_lines.append("-" * 60)
            report_lines.append(f"Current Balance: PHP{Decimal(str(current_balance)):,.2f}".center(60))
            report_lines.append(f"Printed on: {datetime.now().strftime('%d/%m/%Y %H:%M')}".center(60))
            report_lines.append("=" * 60)
            
            full_report = "\n".join(report_lines)
            
            # Show print dialog
            self.show_passbook_print_dialog(full_report, account_num)
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error generating passbook: {e}")
        except ValueError as e:
            messagebox.showerror("Format Error", f"Error formatting transaction data: {e}")

    def print_account_ledger(self):
        account_num = self.ledger_account_entry.get().strip().zfill(4)
        if not account_num or len(account_num) != 4:
            messagebox.showwarning("Error", "Please enter a valid 4-digit account number")
            return
        
        try:
            # Use separate connections for each database
            main_cursor = self.conn.cursor()
            ledger_cursor = self.savings_ledger.cursor()
            
            # Get account info from main DB
            main_cursor.execute("""
                SELECT last_name, first_name 
                FROM accounts 
                WHERE account_number = ?
            """, (account_num,))
            account_info = main_cursor.fetchone()
            
            if not account_info:
                messagebox.showerror("Error", "Account not found")
                return
            
            last_name, first_name = account_info
            
            # Get balance from ledger DB
            ledger_cursor.execute("""
                SELECT current_balance 
                FROM account_balances 
                WHERE account_number = ?
            """, (account_num,))
            balance_result = ledger_cursor.fetchone()
            current_balance = Decimal(str(balance_result[0])) if balance_result else Decimal('0.0')
            
            # Get transactions from ledger DB
            ledger_cursor.execute("""
                SELECT transaction_date, transaction_type, transaction_id, 
                    amount, balance, description
                FROM savings_ledger
                WHERE account_number = ?
                ORDER BY transaction_date
            """, (account_num,))
            
            transactions = ledger_cursor.fetchall()
            
            if not transactions:
                messagebox.showwarning("Error", "No transactions found for this account")
                return
            
            # Create ledger report
            report_lines = []
            
            # Ledger header
            report_lines.append("=" * 90)
            report_lines.append(f"LEDGER FOR ACCOUNT: {account_num}".center(90))
            report_lines.append(f"NAME: {last_name}, {first_name}".center(90))
            report_lines.append("=" * 90)
            report_lines.append(
                "DATE".ljust(15) + 
                "REF NO".ljust(10) + 
                "WITHDRAWAL".ljust(15) + 
                "DEPOSIT".ljust(15) + 
                "INTEREST".ljust(15) + 
                "BALANCE".ljust(15) + 
                "INITIAL".ljust(10)
            )
            report_lines.append("-" * 90)
            
            # Add transactions
            for trans in transactions:
                date = datetime.strptime(trans[0], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%y")
                trans_type = trans[1]
                
                # FIXED: Handle both string and numeric transaction IDs
                try:
                    ref_no = f"{int(trans[2]):05d}"  # Try to convert to integer first
                except (ValueError, TypeError):
                    ref_no = str(trans[2]).zfill(5)  # Fallback to string with zero-padding
                
                amount = Decimal(str(trans[3]))
                balance = Decimal(str(trans[4]))
                
                # Initialize all columns
                withdrawal = ""
                deposit = ""
                interest = ""
                
                # Determine transaction type
                if trans_type in ("Withdrawal", "Debit (Withdrawal)"):
                    withdrawal = f"PHP{amount:,.2f}"
                elif trans_type == "Interest Credit":
                    interest = f"PHP{amount:,.2f}"
                else:  # Deposit, Cheque Deposit, etc.
                    deposit = f"PHP{amount:,.2f}"
                
                report_lines.append(
                    date.ljust(15) + 
                    ref_no.ljust(10) + 
                    withdrawal.ljust(15) + 
                    deposit.ljust(15) + 
                    interest.ljust(15) + 
                    f"PHP{balance:,.2f}".ljust(15) + 
                    "".ljust(10)  # Empty space for initial
                )
            
            # Add footer with current balance
            report_lines.append("-" * 90)
            report_lines.append(f"Current Balance: PHP{Decimal(str(current_balance)):,.2f}".center(90))
            report_lines.append(f"Printed on: {datetime.now().strftime('%d/%m/%Y %H:%M')}".center(90))
            report_lines.append("=" * 90)
            
            full_report = "\n".join(report_lines)
            
            # Show print dialog
            self.show_ledger_print_dialog(full_report, account_num)
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error generating ledger: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {str(e)}")

    def show_passbook_print_dialog(self, report_text, account_num):
        """Dialog for passbook printing with preview"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Passbook Printing")
        dialog.geometry("700x500")
        
        # Preview frame
        preview_frame = ttk.Frame(dialog)
        preview_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Text widget for preview (monospaced font for alignment)
        text = tk.Text(preview_frame, wrap='none', font=('Courier New', 10), width=80, height=25)
        text.insert('1.0', report_text)
        text.config(state='disabled')
        
        scroll_y = ttk.Scrollbar(preview_frame, orient='vertical', command=text.yview)
        scroll_x = ttk.Scrollbar(preview_frame, orient='horizontal', command=text.xview)
        text.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        text.grid(row=0, column=0, sticky='nsew')
        scroll_y.grid(row=0, column=1, sticky='ns')
        scroll_x.grid(row=1, column=0, sticky='ew')
        preview_frame.grid_rowconfigure(0, weight=1)
        preview_frame.grid_columnconfigure(0, weight=1)
        
        # Printing options
        options_frame = ttk.Frame(dialog)
        options_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Button(options_frame, text="Print Passbook", 
                  command=lambda: self.print_passbook(report_text)).pack(side='left', padx=5)
        ttk.Button(options_frame, text="Save as Text", 
                  command=lambda: self.save_passbook_text(report_text, account_num)).pack(side='left', padx=5)
        ttk.Button(options_frame, text="Cancel", 
                  command=dialog.destroy).pack(side='right', padx=5)

    def show_ledger_print_dialog(self, report_text, account_num):
        """Dialog for ledger printing with preview"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Ledger Printing")
        dialog.geometry("800x600")
        
        # Preview frame
        preview_frame = ttk.Frame(dialog)
        preview_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Text widget for preview (monospaced font for alignment)
        text = tk.Text(preview_frame, wrap='none', font=('Courier New', 9), width=100, height=30)
        text.insert('1.0', report_text)
        text.config(state='disabled')
        
        scroll_y = ttk.Scrollbar(preview_frame, orient='vertical', command=text.yview)
        scroll_x = ttk.Scrollbar(preview_frame, orient='horizontal', command=text.xview)
        text.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        text.grid(row=0, column=0, sticky='nsew')
        scroll_y.grid(row=0, column=1, sticky='ns')
        scroll_x.grid(row=1, column=0, sticky='ew')
        preview_frame.grid_rowconfigure(0, weight=1)
        preview_frame.grid_columnconfigure(0, weight=1)
        
        # Printing options
        options_frame = ttk.Frame(dialog)
        options_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Button(options_frame, text="Print Ledger", 
                  command=lambda: self.print_ledger(report_text)).pack(side='left', padx=5)
        ttk.Button(options_frame, text="Save as Text", 
                  command=lambda: self.save_ledger_text(report_text, account_num)).pack(side='left', padx=5)
        ttk.Button(options_frame, text="Cancel", 
                  command=dialog.destroy).pack(side='right', padx=5)

    def print_passbook(self, text):
        """Print the passbook-formatted text"""
        try:
            # Create a temporary file
            temp_file = "passbook_print.txt"
            with open(temp_file, 'w') as f:
                f.write(text)
            
            # For Windows
            if os.name == 'nt':
                os.startfile(temp_file, 'print')
            else:
                # For other platforms, you might need a different approach
                messagebox.showinfo("Printing", "Please print the generated text file")
            
            messagebox.showinfo("Success", "Passbook sent to printer")
        except Exception as e:
            messagebox.showerror("Print Error", f"Could not print passbook: {e}")

    def print_ledger(self, text):
        """Print the ledger-formatted text"""
        try:
            # Create a temporary file
            temp_file = "ledger_print.txt"
            with open(temp_file, 'w') as f:
                f.write(text)
            
            # For Windows
            if os.name == 'nt':
                os.startfile(temp_file, 'print')
            else:
                # For other platforms, you might need a different approach
                messagebox.showinfo("Printing", "Please print the generated text file")
            
            messagebox.showinfo("Success", "Ledger sent to printer")
        except Exception as e:
            messagebox.showerror("Print Error", f"Could not print ledger: {e}")

    def save_passbook_text(self, text, account_num):
        """Save passbook-formatted text to file"""
        filename = f"passbook_{account_num}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
        try:
            with open(filename, 'w') as f:
                f.write(text)
            messagebox.showinfo("Success", f"Passbook saved as {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {e}")

    def save_ledger_text(self, text, account_num):
        """Save ledger-formatted text to file"""
        filename = f"ledger_{account_num}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
        try:
            with open(filename, 'w') as f:
                f.write(text)
            messagebox.showinfo("Success", f"Ledger saved as {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {e}")

    def create_interest_calculation_tab(self):
        """Create the interest calculation tab with proper quarterly compounding"""
        self.interest_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.interest_tab, text="Interest Calculation")
        
        # Create tables if they don't exist
        self.create_quarterly_tables()
        
        # Main frame
        main_frame = ttk.Frame(self.interest_tab, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        # Info section
        info_frame = ttk.LabelFrame(main_frame, text="Quarterly Interest Information", padding=10)
        info_frame.pack(fill='x', pady=5)
        
        ttk.Label(info_frame, 
                text="3% Annual Interest | Compounded Quarterly | Based on Daily Balances",
                font=('Arial', 10, 'bold')).pack()
        
        # Current quarter info
        self.quarter_info = ttk.Label(info_frame, text="", font=('Arial', 9))
        self.quarter_info.pack()
        self.update_quarter_info()
        
        # Calculation controls
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=10)
        
        ttk.Button(control_frame, text="Calculate Interest", 
                command=self.calculate_quarterly_interest).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Apply Interest", 
                command=self.apply_quarterly_interest).pack(side='left', padx=5)
        
        # Results display
        self.results_label = ttk.Label(main_frame, text="", font=('Arial', 10))
        self.results_label.pack()
        
        # Preview frame
        preview_frame = ttk.LabelFrame(main_frame, text="Interest Calculation Preview", padding=10)
        preview_frame.pack(fill='both', expand=True)
        
        # Treeview for displaying calculations
        columns = ("account", "name", "avg_balance", "days", "interest")
        self.interest_tree = ttk.Treeview(preview_frame, columns=columns, show='headings')
        
        # Define headings
        self.interest_tree.heading("account", text="Account #")
        self.interest_tree.heading("name", text="Account Name")
        self.interest_tree.heading("avg_balance", text="Avg Daily Balance")
        self.interest_tree.heading("days", text="Days")
        self.interest_tree.heading("interest", text="Interest")
        
        # Add scrollbars
        scroll_y = ttk.Scrollbar(preview_frame, orient='vertical', command=self.interest_tree.yview)
        scroll_x = ttk.Scrollbar(preview_frame, orient='horizontal', command=self.interest_tree.xview)
        self.interest_tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Layout
        self.interest_tree.pack(side='left', fill='both', expand=True)
        scroll_y.pack(side='right', fill='y')
        scroll_x.pack(side='bottom', fill='x')

    def create_quarterly_tables(self):
        """Create tables for quarterly interest tracking"""
        try:
            cursor = self.conn.cursor()
            
            # Quarterly periods table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS interest_periods (
                    period_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    start_date TEXT NOT NULL,
                    end_date TEXT NOT NULL,
                    is_applied INTEGER DEFAULT 0,
                    applied_date TEXT
                )
            """)
            
            # Quarterly interest calculations
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS quarterly_interest (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    period_id INTEGER NOT NULL,
                    account_number TEXT NOT NULL,
                    avg_daily_balance REAL NOT NULL,
                    days_in_period INTEGER NOT NULL,
                    interest REAL NOT NULL,
                    is_applied INTEGER DEFAULT 0,
                    FOREIGN KEY(period_id) REFERENCES interest_periods(period_id),
                    FOREIGN KEY(account_number) REFERENCES accounts(account_number)
                )
            """)
            
            self.conn.commit()
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error creating quarterly tables: {e}")

    def update_quarter_info(self):
        """Update the current quarter information display"""
        try:
            cursor = self.conn.cursor()
            
            # Get the last period end date
            cursor.execute("SELECT MAX(end_date) FROM interest_periods")
            last_end_date = cursor.fetchone()[0]
            
            if last_end_date:
                last_end = datetime.strptime(last_end_date, "%Y-%m-%d").date()
                days_since_last = (datetime.now().date() - last_end).days
                self.quarter_info.config(
                    text=f"Last quarter ended: {last_end_date} ({days_since_last} days ago)\n"
                        f"Next quarter starts: {last_end + timedelta(days=91)}"
                )
            else:
                # First time setup - assume current quarter started at beginning of current quarter
                today = datetime.now().date()
                quarter_start = datetime(today.year, ((today.month - 1) // 3) * 3 + 1, 1).date()
                days_in_quarter = (today - quarter_start).days
                self.quarter_info.config(
                    text=f"First quarter started: {quarter_start} ({days_in_quarter} days ago)"
                )
                
        except Exception as e:
            messagebox.showerror("Error", f"Error updating quarter info: {e}")

    def calculate_quarterly_interest(self):
        """Calculate interest for all accounts based on daily balances"""
        try:
            # Determine the interest period
            cursor = self.conn.cursor()
            
            # Get the last period end date
            cursor.execute("SELECT MAX(end_date) FROM interest_periods WHERE is_applied = 1")
            last_end_date = cursor.fetchone()[0]
            
            if last_end_date:
                period_start = datetime.strptime(last_end_date, "%Y-%m-%d") + timedelta(days=1)
            else:
                # First time calculation - start from account creation dates
                period_start = self.get_oldest_transaction_date()
            
            period_end = datetime.now()
            days_in_period = (period_end.date() - period_start.date()).days
            
            if days_in_period <= 0:
                messagebox.showwarning("Warning", "No days have passed since last interest calculation")
                return
            
            # Create a new period record
            cursor.execute("""
                INSERT INTO interest_periods (start_date, end_date)
                VALUES (?, ?)
            """, (period_start.strftime("%Y-%m-%d"), period_end.strftime("%Y-%m-%d")))
            period_id = cursor.lastrowid
            
            # Get all accounts with activity in this period
            cursor.execute("""
                SELECT DISTINCT account_number 
                FROM teller_transactions
                WHERE date(transaction_date) BETWEEN date(?) AND date(?)
                UNION
                SELECT account_number FROM accounts
            """, (period_start.strftime("%Y-%m-%d"), period_end.strftime("%Y-%m-%d")))
            
            account_numbers = [row[0] for row in cursor.fetchall()]
            
            # Clear the treeview
            for item in self.interest_tree.get_children():
                self.interest_tree.delete(item)
            
            total_interest = Decimal('0.0')
            calculations = []
            
            # Calculate interest for each account
            for account_num in account_numbers:
                # Get account info
                cursor.execute("""
                    SELECT last_name, first_name FROM accounts WHERE account_number = ?
                """, (account_num,))
                account_info = cursor.fetchone()
                name = f"{account_info[0]}, {account_info[1]}" if account_info else "Unknown"
                
                # Calculate average daily balance for the period
                avg_balance, days_with_balance = self.calculate_average_daily_balance(
                    account_num, period_start, period_end)
                
                if avg_balance <= 0:
                    continue
                
                # Calculate interest (3% annual, compounded quarterly)
                # Daily interest rate = (1 + 0.03)^(1/365) - 1
                daily_rate = (Decimal('1.03') ** (Decimal('1')/Decimal('365'))) - Decimal('1')
                interest = avg_balance * daily_rate * Decimal(days_with_balance)
                interest = interest.quantize(Decimal('0.00'))
                
                # Add to treeview
                self.interest_tree.insert("", "end", values=(
                    account_num,
                    name,
                    f"PHP{avg_balance:,.2f}",
                    days_with_balance,
                    f"PHP{interest:,.2f}"
                ))
                
                total_interest += interest
                
                # Store calculation
                calculations.append((
                    period_id,
                    account_num,
                    float(avg_balance),
                    days_with_balance,
                    float(interest)
                ))
            
            # Save calculations to database
            cursor.executemany("""
                INSERT INTO quarterly_interest 
                (period_id, account_number, avg_daily_balance, days_in_period, interest)
                VALUES (?, ?, ?, ?, ?)
            """, calculations)
            
            self.conn.commit()
            
            self.results_label.config(
                text=f"Calculated interest for {len(calculations)} accounts\n"
                    f"Period: {period_start.strftime('%Y-%m-%d')} to {period_end.strftime('%Y-%m-%d')}\n"
                    f"Total interest: PHP{total_interest:,.2f}"
            )
            
        except sqlite3.Error as e:
            self.conn.rollback()
            messagebox.showerror("Database Error", f"Error calculating interest: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {e}")

    def calculate_average_daily_balance(self, account_num, period_start, period_end):
        """Calculate average daily balance for an account during the period"""
        try:
            cursor = self.conn.cursor()
            
            # Get starting balance (balance just before period_start)
            cursor.execute("""
                SELECT balance FROM (
                    SELECT balance, date 
                    FROM savings_ledger 
                    WHERE account_number = ? AND date(transaction_date) < date(?)
                    ORDER BY date(transaction_date) DESC
                    LIMIT 1
                )
            """, (account_num, period_start.strftime("%Y-%m-%d")))
            
            start_balance_result = cursor.fetchone()
            start_balance = Decimal(str(start_balance_result[0])) if start_balance_result else Decimal('0.0')
            
            # Get all transactions during the period
            cursor.execute("""
                SELECT date(transaction_date) as date, amount, 
                    CASE WHEN transaction_type IN ('Deposit', 'Credit (Deposit)', 'Cheque Deposit') 
                            THEN 1 ELSE -1 END as multiplier
                FROM teller_transactions
                WHERE account_number = ? 
                AND date(transaction_date) BETWEEN date(?) AND date(?)
                ORDER BY date(transaction_date)
            """, (account_num, period_start.strftime("%Y-%m-%d"), period_end.strftime("%Y-%m-%d")))
            
            transactions = cursor.fetchall()
            
            # Calculate daily balances
            daily_balances = {}
            current_balance = start_balance
            current_date = period_start.date()
            end_date = period_end.date()
            
            # Initialize all days with starting balance
            delta = timedelta(days=1)
            while current_date <= end_date:
                daily_balances[current_date] = current_balance
                current_date += delta
            
            # Apply transactions to update balances
            for date_str, amount, multiplier in transactions:
                date = datetime.strptime(date_str, "%Y-%m-%d").date()
                amount = Decimal(str(amount)) * Decimal(str(multiplier))
                
                # Update balance from this date forward
                temp_date = date
                while temp_date <= end_date:
                    daily_balances[temp_date] += amount
                    temp_date += delta
            
            # Calculate average daily balance
            total = Decimal('0.0')
            days_with_balance = 0
            
            for date, balance in daily_balances.items():
                if balance > 0:
                    total += balance
                    days_with_balance += 1
            
            if days_with_balance == 0:
                return Decimal('0.0'), 0
            
            avg_balance = total / Decimal(days_with_balance)
            return avg_balance.quantize(Decimal('0.00')), days_with_balance
            
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error calculating daily balances: {e}")
            return Decimal('0.0'), 0

    def apply_quarterly_interest(self):
        """Apply the calculated quarterly interest to all accounts"""
        try:
            cursor = self.conn.cursor()
            
            # Get the most recent unapplied period
            cursor.execute("""
                SELECT period_id, start_date, end_date 
                FROM interest_periods 
                WHERE is_applied = 0
                ORDER BY end_date DESC
                LIMIT 1
            """)
            
            period = cursor.fetchone()
            if not period:
                messagebox.showwarning("Warning", "No unapplied interest calculations found")
                return
            
            period_id, start_date, end_date = period
            
            # Get all calculations for this period
            cursor.execute("""
                SELECT account_number, interest 
                FROM quarterly_interest 
                WHERE period_id = ? AND is_applied = 0
            """, (period_id,))
            
            calculations = cursor.fetchall()
            
            if not calculations:
                messagebox.showwarning("Warning", "No interest calculations to apply")
                return
            
            today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            total_applied = Decimal('0.0')
            
            for account_num, interest in calculations:
                # Generate transaction number
                trans_num = self.generate_transaction_number()
                
                # Update all three account fields
                cursor.execute("""
                    UPDATE accounts 
                    SET lq_total_savings = total_savings,
                        savings_quarterly_interest = ?,
                        total_savings = total_savings + ?
                    WHERE account_number = ?
                """, (float(interest), float(interest), account_num))
                
                # Update ledger
                description = f"Quarterly Interest ({start_date} to {end_date})"
                self.update_savings_ledger(
                    transaction_id=trans_num,
                    account_number=account_num,
                    trans_type="Interest Credit",
                    amount=float(interest),
                    description=description
                )
                
                # Record transaction
                cursor.execute("""
                    INSERT INTO teller_transactions 
                    (transaction_number, transaction_date, transaction_type, 
                    account_number, amount, teller_id)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (trans_num, today, "Interest Credit", account_num, float(interest), "TELLER1"))
                
                total_applied += Decimal(str(interest))
            
            # Mark period and calculations as applied
            cursor.execute("""
                UPDATE interest_periods 
                SET is_applied = 1, applied_date = ?
                WHERE period_id = ?
            """, (today, period_id))
            
            cursor.execute("""
                UPDATE quarterly_interest 
                SET is_applied = 1
                WHERE period_id = ?
            """, (period_id,))
            
            self.conn.commit()
            
            # Refresh display
            for item in self.interest_tree.get_children():
                self.interest_tree.delete(item)
            
            self.results_label.config(
                text=f"Successfully applied interest for period {start_date} to {end_date}\n"
                    f"Accounts updated: {len(calculations)}\n"
                    f"Total interest applied: PHP{total_applied:,.2f}"
            )
            
            messagebox.showinfo("Success", 
                            f"Applied quarterly interest to {len(calculations)} accounts\n"
                            f"Total: PHP{total_applied:,.2f}")
            
        except sqlite3.Error as e:
            self.conn.rollback()
            messagebox.showerror("Database Error", f"Error applying interest: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = TellerApp(root)
    root.mainloop()